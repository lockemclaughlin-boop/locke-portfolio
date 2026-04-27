"""Build the Vidconvert Studios Webflow assignment as a single .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "Vidconvert_Webflow_Assignment.docx"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_quote(doc, text):
    """Indented italic block (used for prompts pasted into ChatGPT)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def add_rich_para(doc, segments, size=11):
    """segments = [(text, {bold?, italic?, mono?}), ...]"""
    p = doc.add_paragraph()
    for text, opts in segments:
        run = p.add_run(text)
        run.bold = opts.get("bold", False)
        run.italic = opts.get("italic", False)
        run.font.name = "Consolas" if opts.get("mono") else "Calibri"
        run.font.size = Pt(size)
    return p


# ─── Build doc ──────────────────────────────────────────────────────────────
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Web Design — Final Project")
tr.font.size = Pt(14)
tr.font.name = "Calibri"
tr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

h = doc.add_heading("Vidconvert Studios — AI-Generated Webflow Site", level=0)
for r in h.runs:
    r.font.name = "Calibri"

byline = doc.add_paragraph()
byline.add_run("Locke McLaughlin   ·   Subject: Vidconvert Studios (performance marketing + video agency)").italic = True

doc.add_paragraph()

# ─── PART 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, "Part 1 — Research, Content Generation & AI Site", level=1)

# ── 1. Inspiration Sources ──
add_heading(doc, "1. Inspiration Sources (10 pts)", level=2)
add_para(
    doc,
    "Five sites whose specific design choices informed how I directed the AI. For each, I focused on one concrete element (layout, color, typography, navigation, or content structure) rather than a general impression.",
)

inspirations = [
    {
        "name": "Locomotive",
        "url": "https://locomotive.ca",
        "note": (
            "What draws me in is the typographic restraint of the homepage hero — oversized display "
            "type carries the entire message with very little supporting copy, and a single bright accent "
            "color does all the visual heavy lifting. It proves that a confident headline + one accent "
            "color can replace a lot of decoration, which is the discipline I want for Vidconvert's hero."
        ),
    },
    {
        "name": "Active Theory",
        "url": "https://activetheory.net",
        "note": (
            "Their case study grid uses asymmetric video tiles with bold tag chips, and motion is used "
            "as a navigation cue (hover-to-play) rather than decoration. The information hierarchy is "
            "clear at a glance — project name, category, and a moving preview — which is exactly the "
            "structure I want for Vidconvert's Work page."
        ),
    },
    {
        "name": "Buzzworthy Studio",
        "url": "https://buzzworthy.studio",
        "note": (
            "Their 'how we work' section uses a click-through, tabbed layout that lets the user control "
            "pace instead of scrolling past. That interaction model treats process as something worth "
            "exploring, not skimming, and it's the right pattern for Vidconvert's 4-phase Process section."
        ),
    },
    {
        "name": "Linear",
        "url": "https://linear.app",
        "note": (
            "Linear's dark-surface, warm off-white text, and tight letter-spacing on display headings "
            "create a premium feel without ever using more than one accent color. Their palette discipline "
            "is what I want to enforce on Vidconvert: restraint is what makes the yellow accent feel valuable."
        ),
    },
    {
        "name": "Stripe Press",
        "url": "https://press.stripe.com",
        "note": (
            "Their author/founder pages prove that personal storytelling — a real portrait, structured "
            "bio pillars, a clear point of view — earns trust faster than a corporate 'About us'. That's "
            "the model I'm using for the Vidconvert founder section."
        ),
    },
]

for i, site in enumerate(inspirations, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {site['name']} — ").bold = True
    p.add_run(site["url"]).italic = True
    add_para(doc, site["note"])

doc.add_paragraph()

# ── 2. Content & Prompt Planning Guide ──
add_heading(doc, "2. Content & Prompt Planning Guide (40 pts)", level=2)
add_para(
    doc,
    "Below are the eight planning prompts I ran through ChatGPT, with the AI's response captured for each. Responses were lightly edited to align with the brand voice and design system I'd already committed to.",
)

# Prompt 1
add_heading(doc, "Prompt 1 — Site, Audience & Brand Voice", level=3)
add_quote(
    doc,
    "I'm building a simple 3-page website about Vidconvert Studios, a performance marketing and video "
    "production agency that helps ambitious DTC and lead-gen brands turn ad spend into compounding growth "
    "through Meta ads, creative, and full-funnel strategy. Who would visit this site and what would they "
    "be looking for? Based on that audience, also suggest 2–3 words that describe the brand voice and "
    "tone the site should have.",
)
add_para(doc, "Answer:", bold=True)
add_para(
    doc,
    "The primary visitor is a founder, head of marketing, or growth lead at a $1M–$50M DTC, "
    "e-commerce, or lead-generation brand who is either (a) running paid social in-house and hitting a "
    "ceiling, or (b) currently working with an agency that feels like a black box. A smaller secondary "
    "audience is operators at earlier-stage startups ($250K–$1M) who want a senior-level audit before "
    "committing to a long-term partner.",
)
add_para(doc, "What they're looking for, in order of priority:")
add_numbered(
    doc,
    [
        "Proof — real case studies, real numbers, real creative they can actually watch.",
        "A clear point of view — they've talked to ten agencies that all sound the same.",
        "A low-friction next step — a free audit or strategy call, not a 6-month retainer pitch.",
        "Confidence the team can ship creative, not just buy media — because creative is now the lever, not targeting.",
    ],
)
add_rich_para(
    doc,
    [
        ("Brand voice — three words: ", {"bold": True}),
        ("Confident, candid, no-nonsense.", {"bold": True, "italic": True}),
        (
            " The site should sound like a senior operator talking to a peer, not a sales deck. Short "
            "sentences. Strong claims backed by specifics. No buzzwords.",
            {},
        ),
    ],
)

# Prompt 2
add_heading(doc, "Prompt 2 — Target Keywords", level=3)
add_quote(
    doc,
    "Based on that audience, suggest 5 to 8 keywords or phrases my content should naturally include. "
    "Explain why each one is relevant.",
)
add_para(doc, "Answer:", bold=True)
add_para(doc, "Eight keywords/phrases, ranked by intent:")
keywords = [
    ("performance marketing agency", "high-intent commercial keyword; what a founder Googles when they've decided to hire help."),
    ("Meta ads", "the core specialty; signals platform expertise (vs. generic 'social media marketing')."),
    ("creative testing", "differentiator term that filters out brands who think 'creative' means 'logo design'; attracts performance-literate buyers."),
    ("full-funnel growth", "positions the agency above one-trick 'ads only' shops; reinforces the 'flywheel, not a campaign' positioning."),
    ("DTC marketing", "vertical anchor; tells e-commerce brands they're in the right place."),
    ("lead generation", "broadens beyond DTC to the second ICP (B2B services, healthcare like Unlock Patients)."),
    ("ad creative production", "captures founders searching for a team that makes the ads, not just runs them."),
    ("growth audit", "bottom-of-funnel, lead-magnet keyword tied directly to the CTA ('Get your free growth audit')."),
]
for kw, why in keywords:
    add_rich_para(
        doc,
        [
            (f"• {kw}", {"bold": True}),
            (f" — {why}", {}),
        ],
    )

# Prompt 3
add_heading(doc, "Prompt 3 — Site Structure (3 pages)", level=3)
add_quote(
    doc,
    "I'm building a 3-page website about Vidconvert Studios. What are the 3 most logical pages and "
    "what sections should each one contain? Keep it simple.",
)
add_para(doc, "Answer:", bold=True)

add_para(doc, "Page 1 — Home (the convince-them-to-keep-reading page)", bold=True)
add_bullets(
    doc,
    [
        "Sticky nav with logo and primary CTA",
        "Hero: bold headline, value prop, two CTAs, hero video",
        "Client logo bar (social proof)",
        "'Our Approach' — 3-card grid (Science / Style / System)",
        "Yellow services banner / marquee",
        "Final CTA",
    ],
)

add_para(doc, "Page 2 — Work (the proof page)", bold=True)
add_bullets(
    doc,
    [
        "Page header: 'Selected Work'",
        "Featured case study — full-width video",
        "Video grid — 6 case studies with tag chips (Commercial / Social / Event / Brand)",
        "Founder testimonials carousel",
        "CTA strip",
    ],
)

add_para(doc, "Page 3 — About (the trust + conversion page)", bold=True)
add_bullets(
    doc,
    [
        "Page header: 'How we work + who we are'",
        "4-phase Process section (Audit → Build → Scale → Hand off)",
        "Founder bio with portrait and 3 capability pillars (Film, Funnels, AI)",
        "Contact CTA — book a strategy call",
    ],
)

# Prompt 4
add_heading(doc, "Prompt 4 — Layout per Page", level=3)
add_quote(
    doc,
    "Now give me a simple layout description for each page — what sections there are and how they are arranged.",
)
add_para(doc, "Answer:", bold=True)

add_para(doc, "Home — single column, full-bleed sections stacked vertically:", bold=True)
add_numbered(
    doc,
    [
        "Sticky nav (logo left, links center, primary CTA right).",
        "Hero: H1 headline left-aligned, two buttons below, asymmetric 60/40 video grid (one large + two stacked) underneath.",
        "Client logo marquee (full-width horizontal scroll).",
        "Approach: section title + intro paragraph (50/50 split), then 3-column card grid below.",
        "Full-width yellow banner with infinite-scroll services marquee.",
        "Final yellow CTA section, centered text + 2 buttons.",
        "Footer: 4-column (brand / studio / services / contact) plus a giant logo wordmark.",
    ],
)

add_para(doc, "Work — single column:", bold=True)
add_numbered(
    doc,
    [
        "Sticky nav.",
        "Page intro: eyebrow + H1 + CTA on right.",
        "Featured video (16:9 full-width).",
        "3-column grid of 6 video cards (collapses to 2-col tablet, 1-col mobile).",
        "Voices: large quote card with carousel controls.",
        "CTA + footer.",
    ],
)

add_para(doc, "About — single column:", bold=True)
add_numbered(
    doc,
    [
        "Sticky nav.",
        "Process section: 2-column layout — clickable phase list left, detail panel right (collapses to 1-col mobile).",
        "Founder section: 2-column — square portrait left (with rotated yellow backdrop), bio + 3 pillar cards right.",
        "CTA + footer.",
    ],
)

# Prompt 5
add_heading(doc, "Prompt 5 — Generate Placeholder Content", level=3)
add_quote(
    doc,
    "Generate placeholder content for all three pages. Use these target keywords naturally in the "
    "headings and body copy: performance marketing agency, Meta ads, creative testing, full-funnel "
    "growth, DTC marketing, lead generation, ad creative production, growth audit.",
)
add_para(doc, "Answer:", bold=True)

# HOME content
add_para(doc, "HOME", bold=True, size=12)
add_para(doc, "Hero", italic=True)
add_bullets(
    doc,
    [
        "H1: Ad spend. Growth. No fluff.",
        "Sub: Performance marketing and video that drives real results. Full-funnel creative, sharp testing, zero wasted spend.",
        "CTAs: Get started / View work",
    ],
)
add_para(doc, "Client bar: Trusted by ambitious teams shipping real products.", italic=True)
add_para(doc, "Approach", italic=True)
add_bullets(
    doc,
    [
        "Eyebrow: Our approach",
        "H2: We sit at the intersection of science & style.",
        "Card 1 — Science: Test, learn, deploy capital. Structured creative testing frameworks. We make small bets every week so we can confidently double down on the winners.",
        "Card 2 — Style: Creative that earns the click. Native-feeling, scroll-stopping ad creative production built for the platform — not retrofitted from a TV spot.",
        "Card 3 — System: Build a flywheel, not a campaign. Full-funnel growth compounds when creative, audiences, and offer are tuned together. We build all three.",
    ],
)
add_para(doc, "Banner marquee: Meta Ads · Creative Production · Landing Pages · Lifecycle · Analytics", italic=True)
add_para(doc, "CTA", italic=True)
add_bullets(
    doc,
    [
        "H2: Ready to kick off your growth journey?",
        "Body: Book a 30-minute call. We'll audit your last 90 days of Meta spend and tell you exactly where the leaks are. Yours to keep.",
    ],
)

# WORK content
add_para(doc, "WORK", bold=True, size=12)
add_bullets(
    doc,
    [
        "Eyebrow: Selected work",
        "H1: Creative that converts.",
        "Intro: Six recent projects across DTC marketing and lead generation. Hover any tile to play.",
        "Tile — Unlock Patients (Commercial): Healthcare lead-gen brand. New creative system + Meta ads rebuild drove a 2.4× lift in qualified booked calls.",
        "Tile — Chi O Make-A-Wish (Event): Documentary-style fundraising film for a national philanthropy event.",
        "Tile — Berry Glaze (Brand): Full launch campaign for a new DTC food brand — brand film + 12-asset creative test matrix.",
        "Voices header: Founders who've been in the trenches with us.",
    ],
)

# ABOUT content
add_para(doc, "ABOUT", bold=True, size=12)
add_bullets(
    doc,
    [
        "Eyebrow: How we work",
        "H1: A process built for boring, repeatable wins.",
        "Intro: No black boxes. Just four phases, every quarter.",
        "Phase 1 — Audit & diagnose (Week 1): We pull the curtain back on your account, creative library, funnel and offer. You get a detailed growth audit — yours to keep, even if we never work together.",
        "Phase 2 — Build the engine (Weeks 2–4): New campaign architecture, a creative testing matrix, and server-side measurement you can trust.",
        "Phase 3 — Scale with discipline (Month 2+): Three ad creative production concepts shipped per week. Weekly profit reviews. Quarterly offer tests.",
        "Phase 4 — Hand back a flywheel (Always): Every system we build is documented and yours. No agency hostage situations.",
        "Founder section H2: Hi, I'm Locke McLaughlin.",
        "Bio: I graduated from the University of Miami in 2026 with a B.S. in Film — four years obsessing over what makes an audience lean in, click, and care. Most performance marketers think in numbers. Most filmmakers think in emotion. I think in both — and I layer AI automations on top to move faster than any single-discipline shop.",
    ],
)

# Prompt 6
add_heading(doc, "Prompt 6 — Color Palette", level=3)
add_quote(
    doc,
    "Suggest a color palette for a website about Vidconvert Studios, a performance marketing + video "
    "agency with a confident, no-BS voice. Give me 4 colors: a primary, a secondary, an accent, and a "
    "neutral background. Include the hex codes and explain why each color fits the subject.",
)
add_para(doc, "Answer:", bold=True)

palette_table = doc.add_table(rows=5, cols=4)
palette_table.style = "Light Grid Accent 1"
hdr = palette_table.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Hex"
hdr[2].text = "Swatch"
hdr[3].text = "Why it fits"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True

palette_rows = [
    ("Primary (text/ink)", "#F4EFE6", "F4EFE6", "Warmer than pure white — feels editorial and premium, reads beautifully on dark surfaces, signals craft."),
    ("Secondary (surface)", "#18170F", "18170F", "Second dark layer for cards. Subtle warm undertone keeps the palette consistent and avoids 'flat black'."),
    ("Accent", "#FFD60A", "FFD60A", "High-energy, optimistic, unmistakably 'performance' — perfect for CTAs and one full-bleed banner. Used sparingly so it stays loud."),
    ("Neutral background", "#0E0D0B", "0E0D0B", "Lets video case studies pop without competing. Confident, cinematic — the right canvas for a video-led agency."),
]
for i, (role, hex_, fill, why) in enumerate(palette_rows, 1):
    row = palette_table.rows[i].cells
    row[0].text = role
    row[1].text = hex_
    row[2].text = ""
    set_cell_shading(row[2], fill)
    row[3].text = why

add_para(
    doc,
    "Optional secondary accents for category chips and illustration: #3D5AFE (blue), #FF5C39 (coral), "
    "#7CFF6B (mint). Used to color-code work categories — never for primary UI.",
)

# Prompt 7
add_heading(doc, "Prompt 7 — Google Fonts Pairing", level=3)
add_quote(
    doc,
    "Suggest a Google Fonts pairing for a Vidconvert Studios website (a confident performance-marketing "
    "agency with bold display type). Give me one font for headings and one for body text. Explain why "
    "they work together and tell me what sizes to use for H1, H2, and body text.",
)
add_para(doc, "Answer:", bold=True)
add_rich_para(
    doc,
    [
        ("Headings — Bricolage Grotesque ", {"bold": True}),
        (
            "(weights 700–800, letter-spacing −0.03em to −0.05em). A modern grotesque with just enough "
            "personality in the curves to feel editorial rather than corporate. At display sizes it "
            "carries serious weight without looking aggressive.",
            {},
        ),
    ],
)
add_rich_para(
    doc,
    [
        ("Body — Space Grotesk ", {"bold": True}),
        (
            "(weights 400–500). Same geometric DNA as Bricolage so they pair cleanly, but slightly more "
            "neutral and technical, which keeps body copy easy to read.",
            {},
        ),
    ],
)
add_rich_para(
    doc,
    [
        ("Eyebrows / labels — JetBrains Mono ", {"bold": True}),
        (
            "(500–700, uppercase, letter-spacing 0.1em+). The monospace adds a tactical, 'engineered' feel "
            "that reinforces the performance-marketing positioning.",
            {},
        ),
    ],
)
add_para(
    doc,
    "Why they work together: all three are humanist-leaning, share a similar x-height, and were designed "
    "with screen rendering in mind. The contrast comes from weight and case, not from competing typefaces — "
    "a classic editorial system.",
)
add_para(doc, "Sizes:", bold=True)
add_bullets(
    doc,
    [
        "H1: clamp(48px, 8vw, 112px) — line-height 0.95–1.05, weight 800, tracking −0.04em",
        "H2: clamp(36px, 5vw, 64px) — line-height 1.05–1.1, weight 700, tracking −0.03em",
        "H3: 24–28px, weight 700",
        "Body: 16–18px, line-height 1.5–1.65, weight 400",
        "Eyebrow / mono labels: 11–13px, uppercase, letter-spacing 0.12em",
    ],
)

# ── 3. Prompt Log ──
add_heading(doc, "3. Prompt Log — Assembled Webflow Prompt (15 pts)", level=2)
add_para(
    doc,
    "The final assembled prompt I pasted into Webflow's AI site builder, drawing every detail from the "
    "answers above:",
)

# The mega prompt as an indented block
prompt_paras = [
    "Build a 3-page website for Vidconvert Studios, a performance marketing and video production agency "
    "for ambitious DTC and lead-gen brands. Audience: founders and marketing leads at $1M–$50M brands. "
    "Brand voice: confident, candid, no fluff — short sentences, strong claims, no buzzwords.",
    "Pages:",
    "1. Home — Cinematic dark hero with a punchy three-beat headline (\"Ad spend. Growth. No fluff.\"), "
    "a one-line subhead about performance marketing and video that drives real results, and two CTAs: "
    "a yellow primary button labeled \"Get started\" and an outlined ghost button labeled \"View work.\" "
    "Use a dark, atmospheric photo of a team working at a laptop as the hero background. Below the hero: "
    "client logo bar, 3-column \"Our Approach\" section (Science / Style / System), full-width yellow "
    "services banner, final CTA.",
    "2. Work — Featured case study video at top, 6-tile video grid of selected work with category tag "
    "chips (Commercial / Social / Event / Brand), founder testimonial carousel.",
    "3. About — 4-phase \"How we work\" process (Audit → Build → Scale → Hand off), founder bio section "
    "with portrait and three capability pillars (Film, Funnels, AI), contact CTA.",
    "Keywords to weave into supporting copy: performance marketing agency, Meta ads, creative testing, "
    "full-funnel growth, DTC marketing, lead generation, ad creative production, growth audit.",
    "Color palette: background #0E0D0B (near-black), surface #18170F, primary text #F4EFE6 (warm cream), "
    "accent #FFD60A (electric yellow). Use the yellow only for the primary CTA, the full-bleed services "
    "banner, and one or two highlights — never as a body color.",
    "Typography: Bold geometric grotesque for headings (Bricolage Grotesque-style, weights 700–800, "
    "tight letter-spacing −0.03em to −0.04em, all-caps optional on the hero). Clean grotesque for body "
    "(Space Grotesk, weight 400). H1 ~88–112px desktop, H2 ~48–56px, body 16–18px.",
    "Style notes: Dark, editorial, cinematic. Square or lightly-rounded buttons (4–8px radius). Hero "
    "headlines should be terse and punchy — three short beats max, separated by periods — not corporate "
    "marketing copy. Buttons read as labels, not sentences (\"Get started,\" \"View work,\" not \"Book a "
    "free strategy call now\").",
]
for text in prompt_paras:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"

# ── 4. Initial Published Site ──
add_heading(doc, "4. Initial Published Site (15 pts)", level=2)
add_para(
    doc,
    "After pasting the assembled prompt above into Webflow's AI site builder and publishing the result "
    "to the free *.webflow.io subdomain, the initial site is live at:",
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run("Published URL: ")
run.bold = True
run.font.size = Pt(11)
url_run = p.add_run("[PASTE YOUR PUBLISHED WEBFLOW URL HERE]")
url_run.font.size = Pt(11)
url_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_para(doc, "Screenshots of the AI-generated site:", bold=True)
add_bullets(
    doc,
    [
        "[INSERT SCREENSHOT — Home page, full length]",
        "[INSERT SCREENSHOT — Work page, full length]",
        "[INSERT SCREENSHOT — About page, full length]",
    ],
)

# Save
doc.save(OUT)
print(f"Wrote {OUT}")
